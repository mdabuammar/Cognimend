"""
Document Processing Module
Breaks down the long upload_document function into smaller, testable components.
"""
import io
import hashlib
import asyncio
import logging
from typing import List, Optional, Tuple, Dict, Any
from dataclasses import dataclass
import tiktoken
import PyPDF2
import docx

logger = logging.getLogger(__name__)


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    text: str
    chunks: List[str]
    file_hash: str
    chunk_count: int


@dataclass
class ExtractionResult:
    """Result of text extraction."""
    success: bool
    text: str
    error: Optional[str] = None


@dataclass
class ChunkingConfig:
    """Configuration for text chunking."""
    chunk_size: int = 512
    overlap: int = 50
    encoding_name: str = "cl100k_base"


class TextExtractor:
    """
    Handles text extraction from various document formats.
    Follows Single Responsibility Principle.
    """
    
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.txt', '.md'}
    
    @staticmethod
    def extract_from_pdf(file_bytes: bytes) -> ExtractionResult:
        """
        Extract text from PDF file bytes.
        
        Args:
            file_bytes: Raw PDF file bytes
            
        Returns:
            ExtractionResult with extracted text or error
        """
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts: List[str] = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
            
            text = "\n".join(text_parts)
            
            if not text.strip():
                return ExtractionResult(
                    success=False,
                    text="",
                    error="No text could be extracted from PDF"
                )
            
            return ExtractionResult(success=True, text=text)
            
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                error=f"PDF extraction failed: {str(e)}"
            )
    
    @staticmethod
    def extract_from_docx(file_bytes: bytes) -> ExtractionResult:
        """
        Extract text from DOCX file bytes.
        
        Args:
            file_bytes: Raw DOCX file bytes
            
        Returns:
            ExtractionResult with extracted text or error
        """
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n".join(paragraphs)
            
            if not text.strip():
                return ExtractionResult(
                    success=False,
                    text="",
                    error="No text found in DOCX document"
                )
            
            return ExtractionResult(success=True, text=text)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                error=f"DOCX extraction failed: {str(e)}"
            )
    
    @staticmethod
    def extract_from_txt(file_bytes: bytes) -> ExtractionResult:
        """
        Extract text from TXT file bytes.
        
        Args:
            file_bytes: Raw TXT file bytes
            
        Returns:
            ExtractionResult with extracted text or error
        """
        try:
            # Try UTF-8 first, then fallback encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    if text.strip():
                        return ExtractionResult(success=True, text=text)
                except UnicodeDecodeError:
                    continue
            
            # Last resort: ignore errors
            text = file_bytes.decode('utf-8', errors='ignore')
            
            if not text.strip():
                return ExtractionResult(
                    success=False,
                    text="",
                    error="No text found in file"
                )
            
            return ExtractionResult(success=True, text=text)
            
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ExtractionResult(
                success=False,
                text="",
                error=f"TXT extraction failed: {str(e)}"
            )
    
    @classmethod
    def extract(cls, filename: str, file_bytes: bytes) -> ExtractionResult:
        """
        Extract text from file based on extension.
        
        Args:
            filename: Original filename (used to determine format)
            file_bytes: Raw file bytes
            
        Returns:
            ExtractionResult with extracted text or error
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return cls.extract_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            return cls.extract_from_docx(file_bytes)
        elif filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
            return cls.extract_from_txt(file_bytes)
        else:
            return ExtractionResult(
                success=False,
                text="",
                error=f"Unsupported file type. Supported: {cls.SUPPORTED_FORMATS}"
            )
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """Check if file format is supported."""
        return any(filename.lower().endswith(ext) for ext in cls.SUPPORTED_FORMATS)


class TextChunker:
    """
    Handles text chunking with token-based splitting.
    Follows Single Responsibility Principle.
    """
    
    def __init__(self, config: Optional[ChunkingConfig] = None, chunk_size: int = None, overlap: int = None):
        """
        Initialize chunker with config.
        
        Args:
            config: Chunking configuration
            chunk_size: Direct chunk size parameter (alternative to config)
            overlap: Direct overlap parameter (alternative to config)
        """
        if chunk_size is not None or overlap is not None:
            self.config = ChunkingConfig(
                chunk_size=chunk_size or 512,
                overlap=overlap or 50
            )
        else:
            self.config = config or ChunkingConfig()
        self._encoding = None
    
    @property
    def encoding(self):
        """Lazy load tiktoken encoding."""
        if self._encoding is None:
            self._encoding = tiktoken.get_encoding(self.config.encoding_name)
        return self._encoding
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks using tiktoken.
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        tokens = self.encoding.encode(text)
        chunks: List[str] = []
        
        chunk_size = self.config.chunk_size
        overlap = self.config.overlap
        
        start = 0
        while start < len(tokens):
            end = start + chunk_size
            chunk_tokens = tokens[start:end]
            chunk_text = self.encoding.decode(chunk_tokens)
            
            if chunk_text.strip():
                chunks.append(chunk_text)
            
            # Move start forward, but maintain overlap
            start = end - overlap
            
            # Prevent infinite loop for very small texts
            if start <= 0 and end >= len(tokens):
                break
        
        return chunks
    
    def chunk(self, text: str) -> List[str]:
        """
        Alias for chunk_text for API compatibility.
        
        Args:
            text: Text to split
            
        Returns:
            List of text chunks
        """
        return self.chunk_text(text)
    
    def estimate_chunk_count(self, text: str) -> int:
        """
        Estimate number of chunks without actually chunking.
        
        Args:
            text: Text to estimate
            
        Returns:
            Estimated chunk count
        """
        tokens = self.encoding.encode(text)
        effective_chunk_size = self.config.chunk_size - self.config.overlap
        
        if effective_chunk_size <= 0:
            return 1
        
        return max(1, (len(tokens) + effective_chunk_size - 1) // effective_chunk_size)


class FileHasher:
    """Handles file hashing for idempotency checks."""
    
    @staticmethod
    def compute_hash(file_bytes: bytes, algorithm: str = "sha256") -> str:
        """
        Compute hash of file bytes.
        
        Args:
            file_bytes: Raw file bytes
            algorithm: Hash algorithm (sha256, md5, etc.)
            
        Returns:
            Hex digest of hash
        """
        if algorithm == "sha256":
            return hashlib.sha256(file_bytes).hexdigest()
        elif algorithm == "md5":
            return hashlib.md5(file_bytes).hexdigest()
        else:
            raise ValueError(f"Unsupported hash algorithm: {algorithm}")
    
    @staticmethod
    def compute_text_hash(text: str, algorithm: str = "md5") -> str:
        """
        Compute hash of text content.
        
        Args:
            text: Text content
            algorithm: Hash algorithm
            
        Returns:
            Hex digest of hash
        """
        return FileHasher.compute_hash(text.encode('utf-8'), algorithm)


class DocumentProcessor:
    """
    Main document processing orchestrator.
    Combines extraction, chunking, and hashing.
    """
    
    def __init__(
        self,
        chunking_config: Optional[ChunkingConfig] = None,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ):
        """
        Initialize processor.
        
        Args:
            chunking_config: Optional chunking configuration
            chunk_size: Optional chunk size (alternative to chunking_config)
            chunk_overlap: Optional chunk overlap (alternative to chunking_config)
        """
        self.extractor = TextExtractor()
        
        # Support both chunking_config and chunk_size/chunk_overlap parameters
        if chunk_size is not None or chunk_overlap is not None:
            config = ChunkingConfig(
                chunk_size=chunk_size or 500,
                overlap=chunk_overlap or 50
            )
            self.chunker = TextChunker(config)
        else:
            self.chunker = TextChunker(chunking_config)
        
        self.hasher = FileHasher()
    
    def process(
        self,
        filename: str,
        file_bytes: bytes
    ) -> Tuple[bool, Optional[ProcessedDocument], Optional[str]]:
        """
        Process a document completely.
        
        Args:
            filename: Original filename
            file_bytes: Raw file bytes
            
        Returns:
            Tuple of (success, ProcessedDocument or None, error message or None)
        """
        # Compute file hash first
        file_hash = self.hasher.compute_hash(file_bytes)
        
        # Extract text
        extraction_result = self.extractor.extract(filename, file_bytes)
        
        if not extraction_result.success:
            return False, None, extraction_result.error
        
        # Chunk text
        chunks = self.chunker.chunk_text(extraction_result.text)
        
        if not chunks:
            return False, None, "No content could be extracted from document"
        
        return True, ProcessedDocument(
            text=extraction_result.text,
            chunks=chunks,
            file_hash=file_hash,
            chunk_count=len(chunks)
        ), None
    
    def validate_file(self, filename: str, file_bytes: bytes) -> Tuple[bool, Optional[str]]:
        """
        Validate file before processing.
        
        Args:
            filename: Filename to validate
            file_bytes: File content
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        if not filename:
            return False, "Filename is required"
        
        if not file_bytes:
            return False, "File is empty"
        
        if not self.extractor.is_supported(filename):
            return False, f"Unsupported file type. Supported: {TextExtractor.SUPPORTED_FORMATS}"
        
        # Check file size (max 50MB)
        max_size = 50 * 1024 * 1024
        if len(file_bytes) > max_size:
            return False, f"File too large. Maximum size: {max_size // (1024*1024)}MB"
        
        return True, None


class EmbeddingProcessor:
    """
    Handles embedding generation with batching and caching.
    """
    
    def __init__(
        self,
        embedding_client: Optional[Any] = None,
        cache: Optional[Any] = None,
        batch_size: int = 10
    ):
        """
        Initialize embedding processor.
        
        Args:
            embedding_client: Client for generating embeddings
            cache: Optional cache for embeddings
            batch_size: Batch size for processing
        """
        self.client = embedding_client
        self.cache = cache
        self.batch_size = batch_size
    
    async def get_embedding(self, text: str) -> List[float]:
        """
        Get embedding for single text.
        
        Args:
            text: Text to embed
            
        Returns:
            Embedding vector
        """
        # Check cache first
        if self.cache:
            cache_key = f"emb:{hashlib.md5(text.encode()).hexdigest()}"
            cached = await self.cache.get(cache_key)
            if cached:
                return cached
        
        # Get from client or generate mock
        if self.client:
            try:
                embedding = await self.client.get_embedding(text)
            except Exception as e:
                logger.warning(f"Embedding error: {e}, using mock")
                embedding = self._generate_mock_embedding(text)
        else:
            embedding = self._generate_mock_embedding(text)
        
        # Cache result
        if self.cache:
            await self.cache.set(cache_key, embedding, ttl_seconds=86400)
        
        return embedding
    
    async def get_embeddings_batch(
        self,
        texts: List[str]
    ) -> List[List[float]]:
        """
        Get embeddings for multiple texts with batching.
        
        Args:
            texts: List of texts to embed
            
        Returns:
            List of embedding vectors
        """
        embeddings: List[Optional[List[float]]] = [None] * len(texts)
        uncached_indices: List[int] = []
        uncached_texts: List[str] = []
        
        # Check cache for each text
        for idx, text in enumerate(texts):
            if self.cache:
                cache_key = f"emb:{hashlib.md5(text.encode()).hexdigest()}"
                cached = await self.cache.get(cache_key)
                if cached:
                    embeddings[idx] = cached
                    continue
            
            uncached_indices.append(idx)
            uncached_texts.append(text)
        
        # Process uncached in batches
        if uncached_texts:
            for i in range(0, len(uncached_texts), self.batch_size):
                batch = uncached_texts[i:i + self.batch_size]
                batch_indices = uncached_indices[i:i + self.batch_size]
                
                # Process batch in parallel
                tasks = [self.get_embedding(text) for text in batch]
                batch_embeddings = await asyncio.gather(*tasks, return_exceptions=True)
                
                for j, emb in enumerate(batch_embeddings):
                    idx = batch_indices[j]
                    if isinstance(emb, Exception):
                        logger.warning(f"Batch embedding error: {emb}")
                        embeddings[idx] = self._generate_mock_embedding(batch[j])
                    else:
                        embeddings[idx] = emb
        
        return embeddings  # type: ignore
    
    def _generate_mock_embedding(self, text: str) -> List[float]:
        """Generate deterministic mock embedding for testing."""
        import random
        hash_obj = hashlib.md5(text.encode())
        seed = int(hash_obj.hexdigest(), 16)
        random.seed(seed)
        return [random.random() for _ in range(1536)]
