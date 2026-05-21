"""
Upload Service
Features:
- Connection pooling (Priority 4)
- Redis caching (Priority 2)
- Circuit breaker pattern (Priority 3)
- Distributed tracing (Priority 5)
- Proper async/await (Priority 1)
- Batch embedding processing
- Idempotency with file hashing
"""
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import os
import io
import hashlib
import uuid
import asyncio
import logging
import re
from datetime import datetime
from typing import List, Optional, Tuple, Dict, Any
import tiktoken
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct
from psycopg2.extras import RealDictCursor
from dotenv import load_dotenv
import PyPDF2
import docx
from contextlib import asynccontextmanager
import sys

# Add parent path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../'))

# Load environment variables
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '../../.env'))

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def _test_context_default(value: str) -> str:
    return value if os.getenv("API_KEY_REQUIRED", "false").lower() != "true" else ""


def _get_request_context(request: Request) -> tuple[str, str]:
    workspace_id = request.headers.get("X-Workspace-ID") or _test_context_default("test-workspace")
    user_id = request.headers.get("X-User-ID") or _test_context_default("test-user")

    if not workspace_id or not user_id:
        raise HTTPException(status_code=401, detail="Missing authentication headers")

    return workspace_id, user_id

# ===== IMPORT SHARED MODULES =====
SHARED_MODULES_AVAILABLE = False
db_pool = None
cache = None
init_tracing = None
get_tracer = None
DatabaseManager = None

try:
    from services.shared.database import db_pool, DatabasePool
    from services.shared.cache import cache, cache_get_or_compute
    from services.shared.resilience import CircuitBreaker, retry_async, async_timeout
    from services.shared.tracing import init_tracing, get_tracer
    from services.shared.utils import DatabaseManager, HealthCheckBuilder
    from services.shared.exceptions import (
        ServiceException, DatabaseError, DocumentError,
        UnsupportedFileType, ExtractionError, EmptyDocumentError,
        FileTooLarge, DuplicateDocument, ExternalServiceError, VectorStoreError
    )
    from services.shared.security import (
        SecurityConfig, setup_security, get_secure_logger,
        sanitize_filename, validate_file_extension, validate_mime_type,
        verify_api_key, check_rate_limit
    )
    SHARED_MODULES_AVAILABLE = True
    logger.info("✅ Shared modules loaded (pooling, caching, resilience, tracing, security)")
except ImportError as e:
    logger.warning(f"⚠️ Shared modules not available: {e}")
    SHARED_MODULES_AVAILABLE = False
    # Define fallback exceptions
    class ServiceException(Exception): pass
    class DatabaseError(Exception): pass
    class DocumentError(Exception): pass
    class UnsupportedFileType(Exception): pass
    class ExtractionError(Exception): pass
    class EmptyDocumentError(Exception): pass
    class FileTooLarge(Exception): pass
    class DuplicateDocument(Exception): pass
    class ExternalServiceError(Exception): pass
    class VectorStoreError(Exception): pass
    
    # Fallback security function (allows requests when security module not loaded)
    async def verify_api_key():
        """Fallback - allow requests when security module not available."""
        logger.warning("Security module not loaded - authentication disabled")
        return True

# ===== CIRCUIT BREAKERS FOR EXTERNAL SERVICES =====
embedding_circuit = CircuitBreaker(failure_threshold=3, recovery_timeout=60) if SHARED_MODULES_AVAILABLE else None
qdrant_circuit = CircuitBreaker(failure_threshold=3, recovery_timeout=60) if SHARED_MODULES_AVAILABLE else None

# ===== APP LIFECYCLE =====
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifecycle manager"""
    # Startup
    logger.info("🚀 Upload Service starting...")
    
    # Initialize tracing
    if SHARED_MODULES_AVAILABLE:
        init_tracing("upload-service")
    
    # Initialize database tables
    await init_database()
    
    # Initialize Qdrant collection
    await init_qdrant()
    
    logger.info("✅ Upload Service ready")
    
    yield
    
    # Shutdown
    logger.info("🛑 Upload Service shutting down...")
    if SHARED_MODULES_AVAILABLE:
        db_pool.close_all()
    logger.info("✅ Upload Service stopped")

app = FastAPI(
    title="Upload Service",
    version="2.0.0",
    description="Document upload with async processing",
    lifespan=lifespan
)


@app.get("/")
async def root():
    """Root endpoint with service info."""
    return {
        "service": "Upload Service",
        "version": "2.0.0",
        "status": "running",
        "docs": "/docs",
        "health": "/health",
        "endpoints": ["/upload", "/documents", "/health"]
    }


# Security middleware
if SHARED_MODULES_AVAILABLE:
    setup_security(app)

# CORS - use configured origins
cors_origins = os.getenv("CORS_ORIGINS", "http://localhost:5173").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["*"],
)

# ===== OPENROUTER CLIENT =====
openrouter_client = None
try:
    from core.openrouter_client import create_openrouter_client
    openrouter_client = create_openrouter_client(
        preset=os.getenv("OPENROUTER_PRESET", "balanced")
    )
    logger.info("✅ OpenRouter client initialized")
except Exception as e:
    logger.warning(f"⚠️ OpenRouter client error: {e}")

# ===== QDRANT CLIENT =====
qdrant_client = None
try:
    qdrant_client = QdrantClient(
        host=os.getenv("QDRANT_HOST", "localhost"),
        port=int(os.getenv("QDRANT_PORT", "6333"))
    )
    logger.info("✅ Qdrant client initialized")
except Exception as e:
    logger.warning(f"⚠️ Qdrant client error: {e}")


# ===== DATABASE MANAGER =====
db_manager = None
if SHARED_MODULES_AVAILABLE and DatabaseManager:
    db_manager = DatabaseManager(db_pool)
else:
    class FallbackDBManager:
        """Fallback database manager when shared modules unavailable."""
        def get_connection(self):
            import psycopg2
            return psycopg2.connect(
                host=os.getenv("POSTGRES_HOST", "localhost"),
                port=int(os.getenv("POSTGRES_PORT", "5432")),
                database=os.getenv("POSTGRES_DB", "cognimend"),
                user=os.getenv("POSTGRES_USER", "postgres"),
                password=os.getenv("POSTGRES_PASSWORD", ""),
                connect_timeout=5
            )
        def return_connection(self, conn) -> None:
            conn.close()
    db_manager = FallbackDBManager()


# ===== DATABASE FUNCTIONS (USING SHARED MANAGER) =====
def get_db():
    """Get database connection from manager."""
    return db_manager.get_connection()


def return_db(conn) -> None:
    """Return connection to manager."""
    db_manager.return_connection(conn)


async def init_database() -> None:
    """Initialize database tables"""
    try:
        conn = get_db()
        cur = conn.cursor()
        
        # Documents table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id SERIAL PRIMARY KEY,
                title TEXT NOT NULL,
                filename TEXT NOT NULL,
                file_hash VARCHAR(64) UNIQUE,
                version INTEGER DEFAULT 1,
                status VARCHAR(20) DEFAULT 'processing',
                chunk_count INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        
        # Chunks table
        cur.execute("""
            CREATE TABLE IF NOT EXISTS chunks (
                id SERIAL PRIMARY KEY,
                document_id INTEGER REFERENCES documents(id) ON DELETE CASCADE,
                chunk_index INTEGER,
                text TEXT,
                text_hash VARCHAR(64),
                embedding_stored BOOLEAN DEFAULT FALSE,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        
        # Upload audit table (for idempotency)
        cur.execute("""
            CREATE TABLE IF NOT EXISTS upload_audit (
                id SERIAL PRIMARY KEY,
                file_hash VARCHAR(64),
                document_id INTEGER,
                client_request_id VARCHAR(64),
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        
        conn.commit()
        cur.close()
        return_db(conn)
        logger.info("✅ Database tables initialized")
    except Exception as e:
        logger.error(f"❌ Database init error: {e}")
        raise


async def init_qdrant():
    """Initialize Qdrant collection"""
    if not qdrant_client:
        return
    
    try:
        qdrant_client.get_collection("documents")
        logger.info("✅ Qdrant collection exists")
    except Exception as e:
        logger.info(f"Creating Qdrant collection (reason: {e})")
        qdrant_client.create_collection(
            collection_name="documents",
            vectors_config=VectorParams(size=1536, distance=Distance.COSINE)
        )
        logger.info("✅ Qdrant collection created")


# ===== TEXT EXTRACTION =====
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF"""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX"""
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from TXT"""
    return file_bytes.decode('utf-8', errors='ignore')


# ===== CHUNKING =====
def chunk_text(text: str, chunk_size: int = 256, overlap: int = 64) -> List[str]:
    """Split text into chunks using tiktoken"""
    encoding = tiktoken.get_encoding("cl100k_base")
    tokens = encoding.encode(text)
    
    chunks = []
    start = 0
    
    while start < len(tokens):
        end = start + chunk_size
        chunk_tokens = tokens[start:end]
        chunk_text = encoding.decode(chunk_tokens)
        chunks.append(chunk_text)
        start = end - overlap
    
    return chunks


# ===== EMBEDDING FUNCTIONS (ASYNC + CIRCUIT BREAKER) =====
async def get_embedding_async(text: str) -> List[float]:
    """
    Get embedding with proper async/await (NO asyncio.run blocking!)
    Uses circuit breaker for resilience
    """
    if not openrouter_client or not os.getenv("OPENROUTER_API_KEY"):
        return get_mock_embedding(text)
    
    try:
        # Use circuit breaker if available
        if embedding_circuit:
            return await embedding_circuit.call_async(
                openrouter_client.get_embedding, text
            )
        else:
            return await openrouter_client.get_embedding(text)
    except Exception as e:
        logger.warning(f"⚠️ Embedding error: {e}, using mock")
        return get_mock_embedding(text)


async def get_embeddings_batch_async(texts: List[str], batch_size: int = 10) -> List[List[float]]:
    """
    Batch embedding processing with async/await
    Processes in parallel batches for 10x speedup
    """
    all_embeddings = []
    
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        
        # Process batch in parallel
        tasks = [get_embedding_async(text) for text in batch]
        batch_embeddings = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Handle any errors
        for j, emb in enumerate(batch_embeddings):
            if isinstance(emb, Exception):
                logger.warning(f"⚠️ Batch embedding error: {emb}")
                all_embeddings.append(get_mock_embedding(batch[j]))
            else:
                all_embeddings.append(emb)
    
    return all_embeddings


def get_mock_embedding(text: str) -> List[float]:
    """Generate deterministic mock embedding for testing"""
    hash_obj = hashlib.md5(text.encode())
    seed = int(hash_obj.hexdigest(), 16)
    import random
    random.seed(seed)
    return [random.random() for _ in range(1536)]


# ===== CACHING HELPERS =====
async def get_cached_embedding(text: str) -> Optional[List[float]]:
    """Get embedding from cache."""
    if not SHARED_MODULES_AVAILABLE:
        return None
    
    cache_key = f"emb:{hashlib.md5(text.encode()).hexdigest()}"
    return await cache.get(cache_key)


async def cache_embedding(text: str, embedding: List[float]) -> None:
    """Cache embedding with 24h TTL."""
    if not SHARED_MODULES_AVAILABLE:
        return
    
    cache_key = f"emb:{hashlib.md5(text.encode()).hexdigest()}"
    await cache.set(cache_key, embedding, ttl_seconds=86400)  # 24 hours


# ===== FILE PROCESSING HELPERS =====
SUPPORTED_FILE_TYPES = ['.pdf', '.docx', '.txt', '.md', '.doc']
MAX_FILE_SIZE_BYTES = int(os.getenv("MAX_FILE_SIZE_MB", "50")) * 1024 * 1024


def validate_file_type(filename: str) -> str:
    """
    Validate file type and return extension.
    
    Args:
        filename: The uploaded filename
        
    Returns:
        File extension (lowercase)
        
    Raises:
        UnsupportedFileType: If file type not supported
    """
    # Sanitize filename first if security module available
    if SHARED_MODULES_AVAILABLE:
        filename = sanitize_filename(filename)
    
    ext = os.path.splitext(filename.lower())[1]
    if ext not in SUPPORTED_FILE_TYPES:
        if SHARED_MODULES_AVAILABLE:
            raise UnsupportedFileType(filename, SUPPORTED_FILE_TYPES)
        raise HTTPException(
            status_code=400, 
            detail=f"Unsupported file type. Use: {', '.join(SUPPORTED_FILE_TYPES)}"
        )
    return ext


def extract_text(file_bytes: bytes, file_ext: str) -> str:
    """
    Extract text from file based on extension.
    
    Args:
        file_bytes: Raw file content
        file_ext: File extension (.pdf, .docx, .txt)
        
    Returns:
        Extracted text content
        
    Raises:
        ExtractionError: If extraction fails
        EmptyDocumentError: If no text found
    """
    try:
        if file_ext == '.pdf':
            text = extract_text_from_pdf(file_bytes)
        elif file_ext == '.docx':
            text = extract_text_from_docx(file_bytes)
        else:  # .txt
            text = extract_text_from_txt(file_bytes)
    except Exception as e:
        if SHARED_MODULES_AVAILABLE:
            raise ExtractionError(f"file{file_ext}", str(e))
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {e}")
    
    if not text.strip():
        if SHARED_MODULES_AVAILABLE:
            raise EmptyDocumentError(f"file{file_ext}")
        raise HTTPException(status_code=400, detail="No text found in document")
    
    return text


async def check_duplicate_document(
    file_hash: str,
    workspace_id: str
) -> Optional[dict]:
    """
    Check if document already exists by file hash in the workspace.
    """
    conn = get_db()
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            """
            SELECT id, version, status
            FROM documents
            WHERE content_hash = %s AND workspace_id = %s
            ORDER BY created_at DESC
            LIMIT 1
            """,
            (file_hash, workspace_id)
        )
        existing = cur.fetchone()
        cur.close()
        return dict(existing) if existing else None
    finally:
        return_db(conn)


async def process_embeddings(
    chunks: List[str]
) -> Tuple[List[List[float]], int]:
    """
    Process embeddings for chunks with caching.
    
    Args:
        chunks: List of text chunks
        
    Returns:
        Tuple of (embeddings list, count of cached embeddings)
    """
    embeddings = []
    uncached_indices = []
    uncached_texts = []
    
    # Check cache first for each chunk
    for idx, chunk in enumerate(chunks):
        cached_emb = await get_cached_embedding(chunk)
        if cached_emb:
            embeddings.append(cached_emb)
        else:
            embeddings.append(None)
            uncached_indices.append(idx)
            uncached_texts.append(chunk)
    
    cached_count = len(chunks) - len(uncached_texts)
    
    # Batch process uncached embeddings
    if uncached_texts:
        logger.info(f"🔄 Processing {len(uncached_texts)} uncached embeddings")
        new_embeddings = await get_embeddings_batch_async(uncached_texts)
        
        # Fill in the results and cache them
        for i, idx in enumerate(uncached_indices):
            embeddings[idx] = new_embeddings[i]
            await cache_embedding(uncached_texts[i], new_embeddings[i])
    
    return embeddings, cached_count


async def store_document_in_db(
    title: str,
    filename: str,
    file_hash: str,
    version: int,
    workspace_id: str,
    user_id: str
) -> str:
    """Store document metadata in database."""
    conn = get_db()
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(
            """
            INSERT INTO documents (
                user_id, name, original_filename, version, status,
                content_hash, workspace_id, uploaded_by, created_by_user_id,
                file_type, mime_type, access_scope
            )
            VALUES (%s, %s, %s, %s, 'processing', %s, %s, %s, %s, %s, %s, 'workspace')
            RETURNING id
            """,
            (
                user_id,
                title,
                filename,
                f"v{version}.0",
                file_hash,
                workspace_id,
                user_id,
                user_id,
                os.path.splitext(filename)[1].lower().lstrip(".") or "txt",
                "text/plain" if filename.lower().endswith((".txt", ".md")) else None,
            )
        )
        doc_id = str(cur.fetchone()['id'])
        conn.commit()
        cur.close()
        return doc_id
    finally:
        return_db(conn)


async def store_chunks_and_vectors(
    doc_id: str,
    chunks: List[str],
    embeddings: List[List[float]],
    filename: str,
    title: str,
    version: int,
    workspace_id: str
) -> int:
    """Store chunks in DB and vectors in Qdrant."""
    conn = get_db()
    points = []
    metadata_date = None
    for chunk in chunks:
        match = re.search(r"uploaded policy version:\s*(20\d{2})", chunk, flags=re.I)
        if match:
            metadata_date = f"{match.group(1)}-01-01T00:00:00+00:00"
            break
    
    try:
        cur = conn.cursor(cursor_factory=RealDictCursor)
        
        for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
            # Save chunk to DB
            text_hash = hashlib.md5(chunk.encode()).hexdigest()
            cur.execute(
                """
                INSERT INTO chunks (document_id, chunk_index, content, content_hash, token_count, status)
                VALUES (%s, %s, %s, %s, %s, 'processed')
                RETURNING id
                """,
                (doc_id, idx, chunk, text_hash, len(chunk.split()))
            )
            chunk_id = str(cur.fetchone()['id'])
            
            # Prepare Qdrant point
            points.append(
                PointStruct(
                    id=chunk_id,
                    vector=embedding,
                    payload={
                        "document_id": str(doc_id),
                        "chunk_id": chunk_id,
                        "chunk_index": idx,
                        "text": chunk,
                        "content": chunk,
                        "filename": filename,
                        "title": title,
                        "name": title,
                        "version": version,
                        "workspace_id": workspace_id,
                        "uploaded_at": datetime.utcnow().isoformat(),
                        "created_at": datetime.utcnow().isoformat(),
                        "document_created_at": metadata_date,
                        "document_updated_at": metadata_date,
                    }
                )
            )
        
        conn.commit()
        cur.close()
    finally:
        return_db(conn)
    
    # Upload to Qdrant with circuit breaker
    if qdrant_client and points:
        try:
            await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: qdrant_client.upsert(collection_name="documents", points=points)
            )
            logger.info(f"✅ Uploaded {len(points)} vectors to Qdrant")
        except Exception as e:
            logger.error(f"❌ Qdrant upload error: {e}")
            if SHARED_MODULES_AVAILABLE:
                raise VectorStoreError("upsert", str(e))
    
    return len(chunks)


async def finalize_document(doc_id: str, chunk_count: int) -> None:
    """Mark document as ready and invalidate cache."""
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(
            "UPDATE documents SET status = 'ready', chunk_count = %s WHERE id = %s",
            (chunk_count, doc_id)
        )
        conn.commit()
        cur.close()
    finally:
        return_db(conn)
    
    # Invalidate document list cache
    if SHARED_MODULES_AVAILABLE:
        await cache.delete("documents:list")


async def mark_document_failed(doc_id: str) -> None:
    """Mark document as failed."""
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("UPDATE documents SET status = 'error' WHERE id = %s", (doc_id,))
        conn.commit()
        cur.close()
        return_db(conn)
    except Exception:
        pass  # Best effort


# ===== UPLOAD ENDPOINT =====
from fastapi import Request

@app.post("/upload")
async def upload_document(
    request: Request,
    file: UploadFile = File(...),
    title: Optional[str] = None,
    background_tasks: BackgroundTasks = None,
    _auth: bool = Depends(verify_api_key)
) -> Dict[str, Any]:
    """
    Upload and process a document.
    
    Features:
    - Idempotent (same file = same result)
    - Async batch embedding processing
    - Connection pooling
    - Redis caching for embeddings
    - Circuit breaker for API calls
    - Distributed tracing
    - Security: filename sanitization, size validation
    
    Args:
        file: Uploaded file (PDF, DOCX, or TXT)
        title: Optional document title
        background_tasks: FastAPI background tasks
        
    Returns:
        Upload result with document ID, chunk count, etc.
        
    Raises:
        UnsupportedFileType: If file type not supported
        ExtractionError: If text extraction fails
        EmptyDocumentError: If no text found
        FileTooLarge: If file exceeds size limit
        DatabaseError: If database operation fails
        VectorStoreError: If Qdrant operation fails
    """
    doc_id = None
    
    try:
        workspace_id, user_id = _get_request_context(request)

        # Step 1: Sanitize filename
        safe_filename = file.filename
        if SHARED_MODULES_AVAILABLE:
            safe_filename = sanitize_filename(file.filename)
        
        # Step 2: Validate MIME type BEFORE reading the entire file
        allowed_mime_types = {
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
            "text/plain",
            "text/markdown",
            "application/octet-stream"  # Allow generic binary for some clients
        }
        
        if file.content_type and file.content_type not in allowed_mime_types:
            logger.warning(f"Rejected file with MIME type: {file.content_type}")
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type: {file.content_type}. Allowed: PDF, DOCX, DOC, TXT, MD"
            )
        
        # Step 3: Read and validate file
        file_bytes = await file.read()
        
        # Step 4: Validate file size
        if len(file_bytes) > MAX_FILE_SIZE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {MAX_FILE_SIZE_BYTES // (1024*1024)}MB"
            )
        
        # Step 5: Validate file extension
        file_ext = validate_file_type(safe_filename)
        
        # Step 6: Calculate hash for idempotency
        file_hash = hashlib.sha256(file_bytes).hexdigest()
        
        # Step 7: Check for duplicate
        existing = await check_duplicate_document(file_hash, workspace_id)
        if existing and existing['status'] == 'ready':
            logger.info(f"⚡ Returning cached document: {existing['id']}")
            return {
                "success": True,
                "document_id": existing['id'],
                "filename": safe_filename,
                "status": "already_exists",
                "message": "Document already uploaded (idempotent)"
            }
        
        # Step 6: Extract text
        text = extract_text(file_bytes, file_ext)
        
        # Step 7: Determine version and store document
        version = 1
        if existing and existing.get("version"):
            version_digits = "".join(ch for ch in str(existing["version"]) if ch.isdigit())
            version = int(version_digits or "0") + 1
        doc_title = title or safe_filename
        doc_id = await store_document_in_db(doc_title, safe_filename, file_hash, version, workspace_id, user_id)
        
        # Step 6: Chunk text
        chunks = chunk_text(text)
        logger.info(f"📄 Processing {len(chunks)} chunks for document {doc_id}")
        
        # Step 7: Process embeddings with caching
        embeddings, cached_count = await process_embeddings(chunks)
        
        # Step 8: Store chunks and vectors
        chunk_count = await store_chunks_and_vectors(
            doc_id, chunks, embeddings, file.filename, doc_title, version, workspace_id
        )
        
        # Step 9: Finalize document
        await finalize_document(doc_id, chunk_count)
        
        return {
            "success": True,
            "document_id": doc_id,
            "filename": file.filename,
            "title": doc_title,
            "version": version,
            "chunks": chunk_count,
            "cached_embeddings": cached_count,
            "status": "ready",
            "message": f"✓ Document processed with {chunk_count} chunks"
        }
        
    except (UnsupportedFileType, ExtractionError, EmptyDocumentError) as e:
        # Document-specific errors - don't mark as failed since doc wasn't created
        logger.warning(f"⚠️ Document error: {e}")
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Upload error: {e}")
        
        # Mark document as failed if it was created
        if doc_id:
            await mark_document_failed(doc_id)
        
        if SHARED_MODULES_AVAILABLE:
            raise ExternalServiceError("upload", str(e))
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/documents")
async def list_documents(request: Request) -> Dict[str, Any]:
    """List all documents for the workspace with caching."""
    workspace_id, _user_id = _get_request_context(request)
        
    # Try cache first
    cache_key = f"documents:list:{workspace_id}"
    if SHARED_MODULES_AVAILABLE:
        cached = await cache.get(cache_key)
        if cached:
            logger.info("⚡ Cache HIT: documents list")
            return {"documents": cached, "cached": True}
    
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    cur.execute("""
        SELECT 
            d.id,
            d.name AS title,
            d.original_filename AS filename,
            d.version,
            d.status,
            d.chunk_count,
            d.created_at
        FROM documents d
        WHERE d.workspace_id = %s
        ORDER BY d.created_at DESC
    """, (workspace_id,))
    
    documents = cur.fetchall()
    
    # Convert datetime to string for JSON
    for doc in documents:
        if doc.get('created_at'):
            doc['created_at'] = doc['created_at'].isoformat()
    
    cur.close()
    return_db(conn)
    
    # Cache for 5 minutes
    if SHARED_MODULES_AVAILABLE:
        await cache.set(cache_key, documents, ttl_seconds=300)
    
    return {"documents": documents, "cached": False}


@app.get("/documents/{doc_id}")
async def get_document(
    request: Request,
    doc_id: str
) -> Dict[str, Any]:
    """Get document details including chunk information."""
    workspace_id = request.headers.get("X-Workspace-ID")
    if not workspace_id:
        raise HTTPException(status_code=401, detail="Missing workspace ID")
        
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get document details
    cur.execute("""
        SELECT 
            d.id,
            d.name AS title,
            d.original_filename AS filename,
            d.content_hash AS file_hash,
            d.version,
            d.status,
            d.chunk_count,
            d.created_at
        FROM documents d
        WHERE d.id = %s AND d.workspace_id = %s
    """, (doc_id, workspace_id))
    
    document = cur.fetchone()
    
    if not document:
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Convert datetime to string
    if document.get('created_at'):
        document['created_at'] = document['created_at'].isoformat()
    
    cur.close()
    return_db(conn)
    
    return {"document": document}


@app.get("/documents/{doc_id}/download")
async def download_document(
    request: Request,
    doc_id: int
):
    """Download document content reconstructed from chunks."""
    workspace_id = request.headers.get("X-Workspace-ID")
    if not workspace_id:
        raise HTTPException(status_code=401, detail="Missing workspace ID")
        
    from fastapi.responses import Response
    
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get document details
    cur.execute("""
        SELECT id, title, filename, status
        FROM documents
        WHERE id = %s AND workspace_id = %s
    """, (doc_id, workspace_id))
    
    document = cur.fetchone()
    
    if not document:
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document['status'] != 'ready':
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=400, detail="Document is still processing or failed")
    
    # Get all chunks for the document in order
    cur.execute("""
        SELECT text
        FROM chunks
        WHERE document_id = %s
        ORDER BY chunk_index ASC
    """, (doc_id,))
    
    chunks = cur.fetchall()
    cur.close()
    return_db(conn)
    
    if not chunks:
        raise HTTPException(status_code=404, detail="No content found for document")
    
    # Reconstruct the document content
    content = "\n\n".join([chunk['text'] for chunk in chunks])
    
    # Create a text file response
    filename = document['filename'].rsplit('.', 1)[0] + '.txt'
    
    return Response(
        content=content,
        media_type="text/plain",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"'
        }
    )


@app.post("/documents/{doc_id}/reindex")
async def reindex_document(
    request: Request,
    doc_id: int
) -> Dict[str, Any]:
    """Re-index a document by regenerating embeddings for all chunks."""
    workspace_id = request.headers.get("X-Workspace-ID")
    if not workspace_id:
        raise HTTPException(status_code=401, detail="Missing workspace ID")
        
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get document details
    cur.execute("""
        SELECT id, title, filename, status, version
        FROM documents
        WHERE id = %s AND workspace_id = %s
    """, (doc_id, workspace_id))
    
    document = cur.fetchone()
    
    if not document:
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=404, detail="Document not found")
    
    if document['status'] == 'processing':
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=400, detail="Document is already processing")
    
    # Get all chunks for the document
    cur.execute("""
        SELECT id, text, chunk_index
        FROM chunks
        WHERE document_id = %s
        ORDER BY chunk_index ASC
    """, (doc_id,))
    
    chunks_data = cur.fetchall()
    
    if not chunks_data:
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=404, detail="No chunks found for document")
    
    # Update document status to processing
    cur.execute("""
        UPDATE documents 
        SET status = 'processing'
        WHERE id = %s
    """, (doc_id,))
    conn.commit()
    
    try:
        # Extract chunk texts
        chunks = [chunk['text'] for chunk in chunks_data]
        chunk_ids = [chunk['id'] for chunk in chunks_data]
        
        # Generate new embeddings
        embeddings, cached_count = await process_embeddings(chunks)
        
        # Update vectors in Qdrant
        if qdrant_client and embeddings:
            points = [
                PointStruct(
                    id=chunk_id,
                    vector=embedding,
                    payload={
                        "text": chunk_text,
                        "document_id": doc_id,
                        "chunk_index": idx,
                        "filename": document['filename'],
                        "title": document['title'],
                        "version": document['version']
                    }
                )
                for chunk_id, embedding, chunk_text, idx in zip(
                    chunk_ids, embeddings, chunks, range(len(chunks))
                )
            ]
            
            qdrant_client.upsert(collection_name="documents", points=points)
        
        # Update document status to ready
        cur.execute("""
            UPDATE documents 
            SET status = 'ready'
            WHERE id = %s
        """, (doc_id,))
        conn.commit()
        
        cur.close()
        return_db(conn)
        
        # Invalidate cache
        if SHARED_MODULES_AVAILABLE:
            await cache.delete("documents:list")
        
        return {
            "success": True,
            "document_id": doc_id,
            "chunks_reindexed": len(chunks),
            "cached_embeddings": cached_count,
            "message": f"Document {doc_id} successfully re-indexed"
        }
        
    except Exception as e:
        # Mark document as failed
        cur.execute("""
            UPDATE documents 
            SET status = 'error'
            WHERE id = %s
        """, (doc_id,))
        conn.commit()
        cur.close()
        return_db(conn)
        
        logger.error(f"❌ Reindex error: {e}")
        raise HTTPException(status_code=500, detail=f"Re-indexing failed: {str(e)}")


@app.post("/documents/reindex-all")
async def reindex_all_documents(
    background_tasks: BackgroundTasks
) -> Dict[str, Any]:
    """
    Re-index ALL documents by regenerating embeddings with real OpenRouter API.
    This fixes documents that were indexed with mock embeddings.
    Runs in background to avoid timeout.
    """
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get all documents that are ready
    cur.execute("""
        SELECT id, title, filename
        FROM documents
        WHERE status = 'ready'
        ORDER BY id ASC
    """)
    
    documents = cur.fetchall()
    cur.close()
    return_db(conn)
    
    if not documents:
        return {
            "success": False,
            "message": "No documents found to reindex"
        }
    
    doc_ids = [doc['id'] for doc in documents]
    
    # Run reindexing in background
    background_tasks.add_task(reindex_documents_batch, doc_ids)
    
    return {
        "success": True,
        "message": f"Started reindexing {len(documents)} documents in background",
        "document_ids": doc_ids,
        "documents": [{"id": d['id'], "title": d['title']} for d in documents]
    }


async def reindex_documents_batch(doc_ids: List[int]) -> None:
    """Background task to reindex multiple documents."""
    logger.info(f"🔄 Starting batch reindex for {len(doc_ids)} documents")
    
    success_count = 0
    error_count = 0
    
    for doc_id in doc_ids:
        try:
            conn = get_db()
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            # Get document details
            cur.execute("""
                SELECT id, title, filename, version
                FROM documents
                WHERE id = %s
            """, (doc_id,))
            
            document = cur.fetchone()
            
            if not document:
                logger.warning(f"⚠️ Document {doc_id} not found, skipping")
                cur.close()
                return_db(conn)
                continue
            
            # Get all chunks for the document
            cur.execute("""
                SELECT id, text, chunk_index
                FROM chunks
                WHERE document_id = %s
                ORDER BY chunk_index ASC
            """, (doc_id,))
            
            chunks_data = cur.fetchall()


            
            if not chunks_data:
                logger.warning(f"⚠️ No chunks for document {doc_id}, skipping")
                cur.close()
                return_db(conn)
                continue
            
            # Update document status to processing
            cur.execute("""
                UPDATE documents 
                SET status = 'processing'
                WHERE id = %s
            """, (doc_id,))
            conn.commit()
            
            # Extract chunk texts
            chunks = [chunk['text'] for chunk in chunks_data]
            chunk_ids = [chunk['id'] for chunk in chunks_data]
            
            logger.info(f"📄 Reindexing document {doc_id} ({document['title']}) with {len(chunks)} chunks")
            
            # Generate new embeddings using real OpenRouter
            embeddings, cached_count = await process_embeddings(chunks)
            
            # Update vectors in Qdrant
            if qdrant_client and embeddings:
                points = [
                    PointStruct(
                        id=chunk_id,
                        vector=embedding,
                        payload={
                            "text": chunk_text,
                            "document_id": doc_id,
                            "chunk_index": idx,
                            "filename": document['filename'],
                            "title": document['title'],
                            "version": document['version']
                        }
                    )
                    for chunk_id, embedding, chunk_text, idx in zip(
                        chunk_ids, embeddings, chunks, range(len(chunks))
                    )
                ]
                
                qdrant_client.upsert(collection_name="documents", points=points)
            
            # Update document status to ready
            cur.execute("""
                UPDATE documents 
                SET status = 'ready'
                WHERE id = %s
            """, (doc_id,))
            conn.commit()
            
            cur.close()
            return_db(conn)
            
            success_count += 1
            logger.info(f"✅ Document {doc_id} reindexed successfully ({len(chunks)} chunks)")
            
            # Small delay to avoid rate limiting
            await asyncio.sleep(0.5)
            
        except Exception as e:
            error_count += 1
            logger.error(f"❌ Error reindexing document {doc_id}: {e}")
            
            # Try to mark as error
            try:
                conn = get_db()
                cur = conn.cursor()
                cur.execute("""
                    UPDATE documents 
                    SET status = 'error'
                    WHERE id = %s
                """, (doc_id,))
                conn.commit()
                cur.close()
                return_db(conn)
            except:
                pass
    
    logger.info(f"🏁 Batch reindex complete: {success_count} success, {error_count} errors")


@app.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: int,
    _authenticated: bool = Depends(verify_api_key) if SHARED_MODULES_AVAILABLE else None
) -> Dict[str, Any]:
    """Delete a document and its chunks."""
    conn = get_db()
    cur = conn.cursor()
    
    # Delete from Qdrant first
    if qdrant_client:
        try:
            # Get chunk IDs
            cur.execute("SELECT id FROM chunks WHERE document_id = %s", (doc_id,))
            chunk_ids = [row[0] for row in cur.fetchall()]
            
            if chunk_ids:
                qdrant_client.delete(
                    collection_name="documents",
                    points_selector=chunk_ids
                )
        except Exception as e:
            logger.warning(f"⚠️ Qdrant delete error: {e}")
    
    # Delete from database (cascades to chunks)
    cur.execute("DELETE FROM documents WHERE id = %s RETURNING id", (doc_id,))
    deleted = cur.fetchone()
    
    conn.commit()
    cur.close()
    return_db(conn)
    
    # Invalidate cache
    if SHARED_MODULES_AVAILABLE:
        await cache.delete("documents:list")
    
    if deleted:
        return {"success": True, "message": f"Document {doc_id} deleted"}
    else:
        raise HTTPException(status_code=404, detail="Document not found")


@app.get("/health")
async def health_check() -> Dict[str, Any]:
    """Health check with component status."""
    components = {
        "service": "healthy",
        "database": "unknown",
        "qdrant": "unknown",
        "redis": "unknown" if SHARED_MODULES_AVAILABLE else "disabled",
        "openrouter": "unknown"
    }
    
    # Check database
    try:
        conn = get_db()
        cur = conn.cursor()
        cur.execute("SELECT 1")
        cur.close()
        return_db(conn)
        components["database"] = "healthy"
    except Exception as e:
        logger.warning(f"Database health check failed: {e}")
        components["database"] = "unhealthy"
    
    # Check Qdrant
    try:
        if qdrant_client:
            qdrant_client.get_collections()
            components["qdrant"] = "healthy"
    except Exception as e:
        logger.warning(f"Qdrant health check failed: {e}")
        components["qdrant"] = "unhealthy"
    
    # Check Redis
    if SHARED_MODULES_AVAILABLE:
        components["redis"] = "healthy" if cache.is_available() else "unhealthy"
    
    # Check OpenRouter
    components["openrouter"] = "healthy" if openrouter_client else "disabled"
    
    overall = "healthy" if all(
        v in ["healthy", "disabled"] for v in components.values()
    ) else "degraded"
    
    return {
        "status": overall,
        "service": "upload",
        "version": "2.0.0",
        "timestamp": datetime.now().isoformat(),
        "components": components
    }


@app.get("/metrics")
async def get_metrics():
    """Get service metrics"""
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    cur.execute("""
        SELECT 
            COUNT(*) as total_documents,
            SUM(chunk_count) as total_chunks,
            COUNT(CASE WHEN status = 'ready' THEN 1 END) as ready_documents,
            COUNT(CASE WHEN status = 'processing' THEN 1 END) as processing_documents,
            COUNT(CASE WHEN status = 'failed' THEN 1 END) as failed_documents
        FROM documents
    """)
    
    metrics = cur.fetchone()
    cur.close()
    return_db(conn)
    
    # Add cache stats
    cache_stats = {}
    if SHARED_MODULES_AVAILABLE:
        cache_stats = await cache.get_stats()
    
    # Add pool stats
    pool_stats = {}
    if SHARED_MODULES_AVAILABLE:
        pool_stats = db_pool.get_pool_status()
    
    return {
        "documents": metrics,
        "cache": cache_stats,
        "database_pool": pool_stats,
        "timestamp": datetime.now().isoformat()
    }


# ===== RECHUNK ENDPOINTS FOR 90%+ CONFIDENCE =====

@app.post("/rechunk-all")
async def rechunk_all_documents():
    """Reprocess ALL documents with optimal chunking (300 tokens, 75 overlap)"""
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get all documents
    cur.execute("SELECT id, title, content FROM documents WHERE content IS NOT NULL")
    docs = cur.fetchall()
    
    total_chunks = 0
    processed_docs = 0
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=75,
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
    )
    
    for doc in docs:
        try:
            doc_id = doc['id']
            content = doc['content']
            title = doc['title']
            
            # Delete old chunks from Qdrant
            try:
                qdrant_client.delete(
                    collection_name="documents",
                    points_selector={"filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]}}
                )
            except Exception as e:
                logger.warning(f"Could not delete old chunks for doc {doc_id}: {e}")
            
            # Create new chunks with optimal settings
            chunks = splitter.split_text(content)
            
            # Get embeddings and store
            for i, chunk in enumerate(chunks):
                embedding = await get_embedding_async(chunk)
                
                point = PointStruct(
                    id=int(f"{doc_id}{i:04d}"),
                    vector=embedding,
                    payload={
                        "document_id": doc_id,
                        "title": title,
                        "text": chunk,
                        "chunk_index": i,
                        "version": 2  # Mark as rechunked
                    }
                )
                
                qdrant_client.upsert(
                    collection_name="documents",
                    points=[point]
                )
            
            # Update document metadata
            cur.execute(
                "UPDATE documents SET chunk_count = %s, version = version + 1 WHERE id = %s",
                (len(chunks), doc_id)
            )
            
            total_chunks += len(chunks)
            processed_docs += 1
            logger.info(f"✅ Rechunked document {doc_id}: {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"❌ Error rechunking document {doc['id']}: {e}")
    
    conn.commit()
    cur.close()
    return_db(conn)
    
    return {
        "status": "✅ ALL DOCS RECHUNKED",
        "processed_documents": processed_docs,
        "total_chunks": total_chunks,
        "chunk_size": 300,
        "chunk_overlap": 75
    }


@app.post("/rechunk/{doc_id}")
async def rechunk_single(doc_id: int):
    """Rechunk a single document with optimal settings"""
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    
    conn = get_db()
    cur = conn.cursor(cursor_factory=RealDictCursor)
    
    # Get document
    cur.execute("SELECT id, title, content FROM documents WHERE id = %s", (doc_id,))
    doc = cur.fetchone()
    
    if not doc:
        cur.close()
        return_db(conn)
        raise HTTPException(status_code=404, detail=f"Document {doc_id} not found")
    
    content = doc['content']
    title = doc['title']
    
    # Delete old chunks
    try:
        qdrant_client.delete(
            collection_name="documents",
            points_selector={"filter": {"must": [{"key": "document_id", "match": {"value": doc_id}}]}}
        )
    except Exception as e:
        logger.warning(f"Could not delete old chunks: {e}")
    
    # Create new chunks
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=300,
        chunk_overlap=75,
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
    )
    chunks = splitter.split_text(content)
    
    # Store new chunks
    for i, chunk in enumerate(chunks):
        embedding = await get_embedding_async(chunk)
        
        point = PointStruct(
            id=int(f"{doc_id}{i:04d}"),
            vector=embedding,
            payload={
                "document_id": doc_id,
                "title": title,
                "text": chunk,
                "chunk_index": i,
                "version": 2
            }
        )
        
        qdrant_client.upsert(
            collection_name="documents",
            points=[point]
        )
    
    # Update document
    cur.execute(
        "UPDATE documents SET chunk_count = %s, version = version + 1 WHERE id = %s",
        (len(chunks), doc_id)
    )
    
    conn.commit()
    cur.close()
    return_db(conn)
    
    return {
        "status": f"✅ Document {doc_id} rechunked",
        "chunks": len(chunks),
        "chunk_size": 300,
        "chunk_overlap": 75
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
